#!/usr/bin/env python3
"""Convert PRODUCT_KNOWLEDGE.md to a formatted Word document."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(6)


def add_rich_text(paragraph, text: str) -> None:
    """Parse **bold**, `code`, and [text](url) inline markers."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`[]+)"
    )
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("[") and "](" in part:
            match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if match:
                label, url = match.groups()
                run = paragraph.add_run(label)
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
                paragraph.add_run(f" ({url})")
        else:
            paragraph.add_run(part)


def parse_table_rows(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    headers = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return headers, rows


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = ""
        add_rich_text(p, header)
        for run in p.runs:
            run.bold = True

    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row):
            if col_idx < len(cells):
                p = cells[col_idx].paragraphs[0]
                p.text = ""
                add_rich_text(p, cell_text)

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()

    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|$", lines[i + 1].strip()):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            headers, rows = parse_table_rows(table_lines)
            add_table(doc, headers, rows)
            i = j
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, text)
            i += 1
            continue

        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped.startswith("- [x] ")
            text = stripped[6:].strip()
            p = doc.add_paragraph(style="List Bullet")
            prefix = "☑ " if checked else "☐ "
            add_rich_text(p, prefix + text)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:].strip())
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "PRODUCT_KNOWLEDGE.md"
    out = root / "docs" / "CIRRUS_Detachment_Planning_Product_Knowledge.docx"
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    convert_md_to_docx(md, out)
